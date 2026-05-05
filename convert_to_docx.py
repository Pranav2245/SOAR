#!/usr/bin/env python3
"""Convert the SOAR Research Paper markdown to a formatted .docx file."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def parse_md_to_docx(md_path, docx_path):
    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Default Style ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # ── Skip mermaid blocks entirely ──
        if line.strip().startswith('```mermaid'):
            # skip until closing ```
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph('[Diagram — see Mermaid source in markdown version]')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            continue

        # ── Code blocks ──
        if line.strip().startswith('```') and not in_code_block:
            in_code_block = True
            code_lines = []
            i += 1
            continue
        if line.strip() == '```' and in_code_block:
            in_code_block = False
            # Add code block as formatted paragraph
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Tables ──
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Parse row
            cells = [c.strip() for c in line.split('|')[1:-1]]
            # Skip separator rows (|:---|:---|)
            if cells and all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            in_table = False
            # Build table
            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            # Clean markdown bold
                            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            clean = re.sub(r'`(.*?)`', r'\1', clean)
                            cell.text = clean
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    run.font.name = 'Times New Roman'

                            # Bold the header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                doc.add_paragraph()  # spacing after table
            # Don't skip current line, fall through

        # ── Horizontal rules ──
        if line.strip() == '---':
            i += 1
            continue

        # ── Empty lines ──
        if not line.strip():
            i += 1
            continue

        # ── Title (# heading) ──
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # ── Section headings (## ) ──
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # ── Subsection headings (### ) ──
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # ── Sub-sub headings (#### ) ──
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # ── Blockquotes (> ) ──
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # ── Numbered list items ──
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            text = m.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            add_formatted_runs(p, text)
            i += 1
            continue

        # ── Bullet list items ──
        if line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, text)
            i += 1
            continue

        # ── Bold-prefixed lines (like **Abstract**) ──
        if line.startswith('**') and '**' in line[2:]:
            p = doc.add_paragraph()
            add_formatted_runs(p, line)
            i += 1
            continue

        # ── Normal paragraph ──
        p = doc.add_paragraph()
        add_formatted_runs(p, line)
        i += 1

    doc.save(docx_path)
    print(f"[+] Successfully saved: {docx_path}")


def add_formatted_runs(paragraph, text):
    """Parse inline markdown (bold, italic, code) and add as formatted runs."""
    # Pattern: **bold**, *italic*, `code`
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


if __name__ == '__main__':
    import sys
    md_file = sys.argv[1] if len(sys.argv) > 1 else '/Users/pranavsharma/.gemini/antigravity/brain/2bf8db16-3164-4880-b2ae-8485376b5de6/artifacts/research_paper.md'
    out_file = '/Users/pranavsharma/Documents/SOAR/SOAR_Research_Paper.docx'
    parse_md_to_docx(md_file, out_file)
